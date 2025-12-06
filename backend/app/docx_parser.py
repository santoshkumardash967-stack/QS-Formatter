"""
DOCX Parser using python-docx for QS-Formatter.

This parser properly identifies:
- Questions: Paragraphs with the main question numId (usually numId=11 with 100 items)
- Options: Paragraphs with other numIds (4 items each, labeled (a), (b), (c), (d))
- Supplementary text: Normal paragraphs without list numbering

Uses Word's list numbering structure to correctly distinguish questions from options,
regardless of text content or length.
"""
import os
import re
import uuid
import base64
from typing import Optional, List, Dict, Any, Tuple
from docx import Document
from docx.oxml.ns import qn
import mammoth
from bs4 import BeautifulSoup, Tag


def check_needs_image(text: str) -> bool:
    """Check if text contains (Image) marker indicating manual image insertion needed."""
    return '(Image)' in text or '(image)' in text.lower()


def detect_question_type(text: str) -> str:
    """Detect question type from text patterns."""
    text_lower = text.lower()
    if 'assertion' in text_lower or 'reason' in text_lower:
        return 'assertion-reason'
    elif 'match' in text_lower or 'matching' in text_lower:
        return 'matching'
    elif 'how many' in text_lower or 'कितने' in text:
        return 'how-many'
    elif any(x in text_lower for x in ['statement', 'statements', 'कथन', 'कथनों']):
        return 'statement-based'
    elif 'figure' in text_lower or 'image' in text_lower or 'diagram' in text_lower:
        return 'figure-based'
    return 'single'


def get_para_numid(para) -> Optional[int]:
    """Get the numId (list numbering ID) of a paragraph, if any."""
    pPr = para._element.pPr
    if pPr is None:
        return None
    numPr = pPr.numPr
    if numPr is None:
        return None
    numId = numPr.numId
    if numId is None:
        return None
    return numId.val


def get_para_text(para) -> str:
    """
    Get all text from a paragraph, including Math (OMML) elements.
    python-docx's para.text doesn't include text from <m:t> math elements.
    """
    from lxml import etree
    
    text_parts = []
    
    # Namespace for math
    MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    
    # Walk through all elements in the paragraph
    for elem in para._element.iter():
        # Regular text runs
        if elem.tag == qn('w:t'):
            if elem.text:
                text_parts.append(elem.text)
        # Math text elements
        elif elem.tag == f'{{{MATH_NS}}}t':
            if elem.text:
                text_parts.append(elem.text)
    
    return ''.join(text_parts).strip()


def extract_images_from_docx(doc, job_id: str, upload_dir: str) -> Dict[str, List[Dict]]:
    """
    Extract all images from DOCX and return a mapping of paragraph index to images.
    Uses mammoth for image extraction as it handles embedded images well.
    """
    # For now, return empty - images will be handled separately via mammoth
    return {}


def parse_docx_with_numbering(file_path: str, job_id: str, upload_dir: str) -> List[Dict]:
    """
    Parse DOCX file using python-docx to properly identify questions and options
    by their Word list numbering.
    
    Two document styles are supported:
    1. Hindi style: Questions use numId=11 (100 items), options use other numIds (4 items each)
    2. English style: Questions have no numId (plain paragraphs), options use numIds (4 items each)
    
    Options always have (a), (b), (c), (d) formatting via Word list numbering.
    """
    doc = Document(file_path)
    
    # First pass: analyze numId distribution
    numid_counts = {}
    for para in doc.paragraphs:
        numid = get_para_numid(para)
        if numid is not None:
            numid_counts[numid] = numid_counts.get(numid, 0) + 1
    
    if not numid_counts:
        print("Warning: No numbered lists found in document")
        return []
    
    # Detect document style
    max_numid = max(numid_counts, key=numid_counts.get)
    max_count = numid_counts[max_numid]
    
    # If the most common numId has ~100 items, questions are numbered (Hindi style)
    # If all numIds have 4 items, questions are unnumbered paragraphs (English style)
    if max_count >= 50:
        # Hindi style: questions have their own numId
        question_numid = max_numid
        print(f"Hindi style: question numId={question_numid} with {max_count} items")
        return parse_hindi_style(doc, question_numid, job_id, upload_dir)
    else:
        # English style: questions are plain paragraphs, all numIds are options
        print(f"English style: all numIds have 4 items (options only)")
        return parse_english_style(doc, numid_counts, job_id, upload_dir)


def parse_hindi_style(doc, question_numid: int, job_id: str, upload_dir: str) -> List[Dict]:
    """Parse Hindi-style document where questions have their own numId."""
    questions = []
    current_question = None
    current_options = []
    current_option_numid = None
    supplementary_text = []
    q_id = 1
    
    for para in doc.paragraphs:
        text = get_para_text(para)
        if not text:
            continue
            
        numid = get_para_numid(para)
        
        if numid == question_numid:
            # This is a question paragraph - save previous question first
            if current_question is not None:
                if len(current_options) == 4:
                    questions.append(finalize_question(
                        q_id, current_question, supplementary_text, 
                        current_options, job_id, upload_dir
                    ))
                    q_id += 1
                else:
                    print(f"Warning: Q{q_id} has {len(current_options)} options, skipping")
                
                current_options = []
                current_option_numid = None
                supplementary_text = []
            
            current_question = text
            
        elif numid is not None and numid != question_numid:
            # This is an option paragraph
            current_options.append({
                'text': text,
                'needs_image': check_needs_image(text)
            })
            current_option_numid = numid
                
        else:
            # Regular paragraph - supplementary text for current question
            if current_question is not None:
                supplementary_text.append(text)
    
    # Save last question
    if current_question is not None and len(current_options) == 4:
        questions.append(finalize_question(
            q_id, current_question, supplementary_text,
            current_options, job_id, upload_dir
        ))
    
    print(f"Parsed {len(questions)} questions (Hindi style)")
    return questions


def parse_english_style(doc, numid_counts: Dict, job_id: str, upload_dir: str) -> List[Dict]:
    """
    Parse English-style document where questions are plain paragraphs
    and options are numbered with (a), (b), (c), (d).
    
    Pattern:
    [Plain paragraphs] = Q1 question text
    [4 numbered items with same numId] = Q1 options (a)(b)(c)(d)
    [Plain paragraphs] = Q2 question text
    [4 numbered items with different numId] = Q2 options
    ... repeat
    
    Key insight: When we see a plain paragraph AFTER having 4 options,
    that's the boundary - save the previous question and start a new one.
    """
    questions = []
    current_question_parts = []
    current_options = []
    q_id = 1
    
    for para in doc.paragraphs:
        text = get_para_text(para)
        if not text:
            continue
            
        numid = get_para_numid(para)
        
        if numid is not None and numid != 0:  # Skip numId=0 (headers/footers)
            # This is an option
            current_options.append({
                'text': text,
                'needs_image': check_needs_image(text)
            })
            
        else:
            # Plain paragraph (question text or supplementary)
            
            # If we already have 4 options, we've completed a question!
            if len(current_options) == 4 and current_question_parts:
                question_text = current_question_parts[0]
                supplementary = current_question_parts[1:] if len(current_question_parts) > 1 else []
                questions.append(finalize_question(
                    q_id, question_text, supplementary,
                    current_options, job_id, upload_dir
                ))
                q_id += 1
                # Reset for next question
                current_question_parts = []
                current_options = []
            
            # If we have partial options (1-3), that's an error - reset
            elif len(current_options) > 0 and len(current_options) < 4:
                print(f"Warning: Q{q_id} had only {len(current_options)} options before new question text")
                current_options = []
            
            # Add this paragraph to current question parts
            current_question_parts.append(text)
    
    # Don't forget the last question!
    if current_question_parts and len(current_options) == 4:
        question_text = current_question_parts[0]
        supplementary = current_question_parts[1:] if len(current_question_parts) > 1 else []
        questions.append(finalize_question(
            q_id, question_text, supplementary,
            current_options, job_id, upload_dir
        ))
    
    print(f"Parsed {len(questions)} questions (English style)")
    return questions


def finalize_question(q_id: int, question_text: str, supplementary: List[str],
                      options: List[Dict], job_id: str, upload_dir: str) -> Dict:
    """Create a finalized question dictionary."""
    
    # Combine question text with supplementary paragraphs
    full_question = question_text
    if supplementary:
        full_question += "\n" + "\n".join(supplementary)
    
    # Build options with labels
    option_labels = ['A', 'B', 'C', 'D']
    formatted_options = []
    for i, opt in enumerate(options):
        formatted_options.append({
            'label': option_labels[i],
            'english_text': opt['text'],
            'hindi_text': '',
            'is_correct': False,
            'images': [],
            'needs_image': opt['needs_image']
        })
    
    # Check flags
    q_needs_image = check_needs_image(full_question)
    flags = []
    if q_needs_image:
        flags.append('needs_image')
    if any(opt['needs_image'] for opt in formatted_options):
        flags.append('options_need_images')
    
    return {
        'id': q_id,
        'english_text': full_question,
        'hindi_text': '',
        'question_type': detect_question_type(full_question),
        'options': formatted_options,
        'answer': '',
        'solution_english': '',
        'solution_hindi': '',
        'grading': '',
        'tables': [],
        'images': [],
        'confidence': 1.0,
        'flags': flags,
        'needs_image': q_needs_image
    }


def parse_document(file_path: str, job_id: str, upload_dir: str) -> List[Dict]:
    """Main entry point - try numId-based parsing first, fall back to mammoth."""
    try:
        questions = parse_docx_with_numbering(file_path, job_id, upload_dir)
        if questions:
            return questions
    except Exception as e:
        print(f"NumId-based parsing failed: {e}, falling back to mammoth parser")
    
    # Fall back to original mammoth-based parser
    from . import smart_parser
    return smart_parser.parse_docx_smart(file_path, job_id, upload_dir)


# Test
if __name__ == "__main__":
    import json
    
    # Test with MOCK files
    en_file = "/workspaces/QS-Formatter/files/MOCK 1 ENGLISH QUESTION.docx"
    hi_file = "/workspaces/QS-Formatter/files/MOCK 1 HINDI QUESTION.docx"
    
    print("=" * 60)
    print("Testing English file...")
    en_questions = parse_docx_with_numbering(en_file, "test", "/tmp")
    print(f"Extracted {len(en_questions)} English questions")
    
    print("\n" + "=" * 60)
    print("Testing Hindi file...")
    hi_questions = parse_docx_with_numbering(hi_file, "test", "/tmp")
    print(f"Extracted {len(hi_questions)} Hindi questions")
    
    # Compare Q64
    if len(en_questions) >= 64 and len(hi_questions) >= 64:
        print("\n" + "=" * 60)
        print("Q64 English:")
        print(f"  {en_questions[63]['english_text'][:80]}")
        for opt in en_questions[63]['options']:
            print(f"  ({opt['label'].lower()}) {opt['english_text'][:40]}")
        
        print("\nQ64 Hindi:")
        print(f"  {hi_questions[63]['english_text'][:80]}")
        for opt in hi_questions[63]['options']:
            print(f"  ({opt['label'].lower()}) {opt['english_text'][:40]}")
