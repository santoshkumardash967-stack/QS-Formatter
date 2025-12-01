"""
Exporter Module - Generates final DOCX in the required format.
"""
import os
from typing import List
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from bs4 import BeautifulSoup

from .models import Question, TableData, ImageData, TableRenderMode


class DOCXExporter:
    """
    Exports questions to DOCX in the strict required format.
    """
    
    def __init__(self, output_dir: str = "/tmp/qs-formatter"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def export(self, questions: List[Question], job_id: str) -> str:
        """
        Export questions to DOCX file.
        Returns the path to the exported file.
        """
        doc = Document()
        
        # Set up document styles
        self._setup_styles(doc)
        
        for i, question in enumerate(questions):
            self._add_question(doc, question)
            
            # Add two blank lines after each question (except the last)
            if i < len(questions) - 1:
                doc.add_paragraph()
                doc.add_paragraph()
        
        # Save document
        output_path = os.path.join(self.output_dir, job_id, f"output_{job_id}.docx")
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return output_path
    
    def _setup_styles(self, doc: Document):
        """Set up document styles for consistent formatting."""
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Noto Sans'
        font.size = Pt(11)
        
        # Set font for Hindi (Devanagari)
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:cs'), 'Noto Sans Devanagari')
    
    def _add_question(self, doc: Document, question: Question):
        """Add a single question to the document."""
        # Q{n}. English question text?
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"Q{question.id}. {question.english_text}")
        q_run.font.name = 'Noto Sans'
        q_run.font.size = Pt(11)
        
        # Hindi question (indented with 4 spaces)
        if question.hindi_text:
            hi_para = doc.add_paragraph()
            hi_run = hi_para.add_run(f"    ({question.hindi_text})")
            hi_run.font.name = 'Noto Sans Devanagari'
            hi_run.font.size = Pt(11)
            self._set_hindi_font(hi_run)
        
        # Add tables if present (before Type line for questions with tables in content)
        for table_data in question.tables:
            self._add_table(doc, table_data)
        
        # Type: line
        type_para = doc.add_paragraph()
        type_run = type_para.add_run(f"Type: {question.question_type.value}")
        type_run.font.name = 'Noto Sans'
        type_run.font.size = Pt(11)
        
        # Options
        for option in question.options:
            opt_para = doc.add_paragraph()
            
            # (A) English option (Hindi option)
            en_text = option.english_text or ""
            hi_text = option.hindi_text or ""
            
            if hi_text:
                opt_text = f"({option.label}) {en_text} ({hi_text})"
            else:
                opt_text = f"({option.label}) {en_text}"
            
            opt_run = opt_para.add_run(opt_text)
            opt_run.font.name = 'Noto Sans'
            opt_run.font.size = Pt(11)
            
            # Apply Hindi font to the Hindi portion
            if hi_text:
                self._set_hindi_font(opt_run)
        
        # Add images if present
        for image_data in question.images:
            self._add_image(doc, image_data)
        
        # Answer line (if present)
        if question.answer:
            ans_para = doc.add_paragraph()
            ans_run = ans_para.add_run(f"Ans. {question.answer}")
            ans_run.font.name = 'Noto Sans'
            ans_run.font.size = Pt(11)
        
        # Solution (if present)
        if question.solution_english:
            sol_para = doc.add_paragraph()
            sol_text = f"Sol: {question.solution_english}"
            if question.solution_hindi:
                sol_text += f"\n({question.solution_hindi})"
            sol_run = sol_para.add_run(sol_text)
            sol_run.font.name = 'Noto Sans'
            sol_run.font.size = Pt(11)
        
        # Grading (if present)
        if question.grading:
            grade_para = doc.add_paragraph()
            grade_run = grade_para.add_run(f"Grading: {question.grading}")
            grade_run.font.name = 'Noto Sans'
            grade_run.font.size = Pt(11)
    
    def _set_hindi_font(self, run):
        """Set Hindi/Devanagari font for a run."""
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(qn('w:cs'), 'Noto Sans Devanagari')
    
    def _add_table(self, doc: Document, table_data: TableData):
        """Add a table to the document."""
        if table_data.render_mode == TableRenderMode.IMAGE and table_data.image_path:
            # Add as image
            if os.path.exists(table_data.image_path):
                doc.add_picture(table_data.image_path, width=Inches(5.5))
        else:
            # Parse HTML table and recreate
            self._recreate_table_from_html(doc, table_data.html)
    
    def _recreate_table_from_html(self, doc: Document, html: str):
        """Recreate a table from HTML."""
        soup = BeautifulSoup(html, 'lxml')
        html_table = soup.find('table')
        
        if not html_table:
            return
        
        rows = html_table.find_all('tr')
        if not rows:
            return
        
        # Determine table dimensions
        max_cols = 0
        for row in rows:
            cells = row.find_all(['td', 'th'])
            max_cols = max(max_cols, len(cells))
        
        if max_cols == 0:
            return
        
        # Create table
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        # Populate table
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    text = cell.get_text(' ', strip=True)
                    table.rows[i].cells[j].text = text
                    
                    # Bold for header cells
                    if cell.name == 'th' or cell.find('strong'):
                        for paragraph in table.rows[i].cells[j].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
    
    def _add_image(self, doc: Document, image_data: ImageData):
        """Add an image to the document."""
        if os.path.exists(image_data.path):
            try:
                doc.add_picture(image_data.path, width=Inches(4))
            except Exception as e:
                # Log error but continue
                print(f"Error adding image {image_data.path}: {e}")


class SimpleExporter:
    """
    Simplified exporter that produces exact format from prompt.
    
    Format:
    Q1. English question text?
        (Hindi question text?)
    Type: multiple_choice
    (A) English option (Hindi option)
    (B) ...
    (C) ...
    (D) ...
    
    [blank line]
    [blank line]
    """
    
    def __init__(self, output_dir: str = "/tmp/qs-formatter"):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
    
    def export(self, questions: List[Question], job_id: str) -> str:
        """Export questions to DOCX in strict format."""
        doc = Document()
        
        # Setup fonts
        self._setup_document(doc)
        
        for i, q in enumerate(questions):
            # Q{n}. English text
            p1 = doc.add_paragraph()
            run1 = p1.add_run(f"Q{q.id}. {q.english_text}")
            self._apply_font(run1)
            
            # (Hindi text) - indented
            if q.hindi_text:
                p2 = doc.add_paragraph()
                run2 = p2.add_run(f"    ({q.hindi_text})")
                self._apply_font(run2, hindi=True)
            
            # Tables (if any, add before Type)
            for table in q.tables:
                self._add_table(doc, table)
            
            # Type: multiple_choice
            p3 = doc.add_paragraph()
            run3 = p3.add_run(f"Type: {q.question_type.value}")
            self._apply_font(run3)
            
            # Options: (A) English (Hindi)
            for opt in q.options:
                p_opt = doc.add_paragraph()
                
                if opt.hindi_text:
                    opt_text = f"({opt.label}) {opt.english_text} ({opt.hindi_text})"
                else:
                    opt_text = f"({opt.label}) {opt.english_text}"
                
                run_opt = p_opt.add_run(opt_text)
                self._apply_font(run_opt, hindi=bool(opt.hindi_text))
            
            # Images (if any)
            for img in q.images:
                self._add_image(doc, img)
            
            # Answer (if present)
            if q.answer:
                p_ans = doc.add_paragraph()
                run_ans = p_ans.add_run(f"Ans. {q.answer}")
                self._apply_font(run_ans)
            
            # Solution (if present)
            if q.solution_english:
                p_sol = doc.add_paragraph()
                sol_text = f"Sol: {q.solution_english}"
                if q.solution_hindi:
                    sol_text += f"\n({q.solution_hindi})"
                run_sol = p_sol.add_run(sol_text)
                self._apply_font(run_sol)
            
            # Grading (if present)
            if q.grading:
                p_grade = doc.add_paragraph()
                run_grade = p_grade.add_run(f"Grading: {q.grading}")
                self._apply_font(run_grade)
            
            # Two blank lines after each question
            doc.add_paragraph()
            doc.add_paragraph()
        
        # Save
        output_dir = os.path.join(self.output_dir, job_id)
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, f"formatted_output.docx")
        doc.save(output_path)
        
        return output_path
    
    def _setup_document(self, doc: Document):
        """Setup document with appropriate fonts."""
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Paragraph spacing
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.15
    
    def _apply_font(self, run, hindi: bool = False):
        """Apply font settings to a run."""
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        
        if hindi:
            # Set complex script font for Hindi
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(qn('w:cs'), 'Noto Sans Devanagari')
    
    def _add_table(self, doc: Document, table_data: TableData):
        """Add table to document."""
        if table_data.render_mode == TableRenderMode.IMAGE and table_data.image_path:
            if os.path.exists(table_data.image_path):
                doc.add_picture(table_data.image_path, width=Inches(5.5))
            return
        
        # Parse and recreate table
        soup = BeautifulSoup(table_data.html, 'lxml')
        html_table = soup.find('table')
        
        if not html_table:
            return
        
        rows = html_table.find_all('tr')
        if not rows:
            return
        
        max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
        if max_cols == 0:
            return
        
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    text = cell.get_text(' ', strip=True)
                    table.rows[i].cells[j].text = text
                    
                    if cell.name == 'th' or cell.find('strong'):
                        for para in table.rows[i].cells[j].paragraphs:
                            for run in para.runs:
                                run.bold = True
    
    def _add_image(self, doc: Document, image_data: ImageData):
        """Add image to document."""
        if os.path.exists(image_data.path):
            try:
                doc.add_picture(image_data.path, width=Inches(4))
            except Exception as e:
                print(f"Error adding image: {e}")
